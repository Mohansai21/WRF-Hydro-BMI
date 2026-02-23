#!/usr/bin/env python3
"""
Generate Doc 13: SCHISM & Its BMI Complete Deep Dive — Word Document
Creates a professional .docx file with tables, formatting, and structure.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._element.get_or_add_tcPr(), qn('w:shd'))
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')

def add_table_row(table, cells_data, bold_first=False, header=False):
    """Add a row to a table with formatting."""
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        if header:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_shading(cell, '2E4057')
        elif bold_first and i == 0:
            run.bold = True

def create_document():
    doc = Document()

    # --- Page setup ---
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # --- Styles ---
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Calibri'
        h_style.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('SCHISM & Its BMI\nComplete Deep Dive')
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Doc 13 — WRF-Hydro BMI Wrapper Project')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x5D, 0x8A, 0xA8)

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('Everything about SCHISM, its physics, its BMI wrapper,\n'
                       'input/output variables, grid system, t0/t1 pattern,\n'
                       'and implementation patterns.\n\n'
                       'Written for ML Engineers with no prior oceanography background.')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run('February 2026')
    run.font.size = Pt(14)
    run.bold = True

    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS
    # =========================================================================
    doc.add_heading('Table of Contents', level=1)

    toc_items = [
        ('PART I: SCHISM — THE MODEL', [
            '1. What is SCHISM?',
            '2. SCHISM Physics & What It Computes',
            '3. The Unstructured Mesh',
            '4. Vertical Coordinate System',
            '5. Semi-Implicit Time Stepping',
            '6. SCHISM vs Other Ocean Models',
            '7. SCHISM in NOAA Operations',
            '8. SCHISM Modules & Capabilities',
        ]),
        ('PART II: SCHISM BMI — THE WRAPPER', [
            '9. What is BMI? (Quick Recap)',
            '10. SCHISM BMI Architecture',
            '11. BMI Initialize',
            '12. BMI Update',
            '13. BMI Finalize',
            '14. BMI Input Variables (12 Detailed)',
            '15. BMI Output Variables (5 Detailed)',
            '16. BMI Grid System (9 Grids)',
            '17. The t0/t1 Sliding Window Pattern',
            '18. RAINRATE — The Special Variable',
            '19. Variable Info Functions',
            '20. Get/Set Value Patterns',
            '21. Time Functions',
            '22. NextGen Integration',
            '23. Build System & Configuration',
            '24. Current Limitations',
            '25. Repository Links & References',
            '26. Summary & Key Takeaways',
        ]),
    ]

    for part_name, items in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(part_name)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)
        for item in items:
            p = doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # =========================================================================
    # PART I: SCHISM — THE MODEL
    # =========================================================================
    part1 = doc.add_heading('PART I: SCHISM — THE MODEL', level=1)

    # --- Section 1: What is SCHISM ---
    doc.add_heading('1. What is SCHISM?', level=2)

    doc.add_heading('Full Name', level=3)
    doc.add_paragraph('Semi-implicit Cross-scale Hydroscience Integrated System Model')

    doc.add_heading('Definition', level=3)
    doc.add_paragraph(
        'SCHISM is an open-source, community-supported, 3D coastal ocean model that simulates '
        'water levels, currents, temperature, salinity, and waves on unstructured meshes — '
        'from tiny creeks to the entire Atlantic Ocean in a single simulation.'
    )

    doc.add_heading('ML Analogy', level=3)
    doc.add_paragraph(
        'Think of SCHISM as a Graph Neural Network (GNN) simulator for water. '
        'Where a CNN operates on a regular pixel grid, SCHISM operates on a graph '
        'of variable-density nodes (triangles and quads). Message passing between '
        'neighbors is the finite element/volume discretization. Multi-scale feature '
        'extraction maps to creek-to-ocean resolution in one mesh.'
    )

    doc.add_heading('Who Built SCHISM?', level=3)
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    add_table_row(t, ['Aspect', 'Details'], header=True)
    rows = [
        ('Institution', 'Virginia Institute of Marine Science (VIMS), College of William & Mary'),
        ('Lead Developer', 'Dr. Y. Joseph Zhang (Professor of Marine Science)'),
        ('Education', 'Ph.D. University of Wollongong, Australia (1996)'),
        ('Contact', 'yjzhang@vims.edu'),
        ('ORCID', '0000-0002-2561-1241'),
        ('License', 'Apache 2.0 (open source)'),
    ]
    for r in rows:
        add_table_row(t, r, bold_first=True)

    doc.add_heading('History & Evolution', level=3)
    doc.add_paragraph(
        '1996-2008: SELFE model developed by Dr. Zhang at Oregon Health & Science University. '
        'Foundational paper: Zhang & Baptista (2008), Ocean Modelling.\n\n'
        '2014: Dr. Zhang moves to VIMS. Forks SELFE v3.1dc with major upgrades.\n\n'
        '2016: SCHISM released. Paper: Zhang, Ye, Stanev, Grashorn (2016), Ocean Modelling. '
        'Key upgrades: LSC2 vertical coordinates, hybrid tri-quad mesh, model polymorphism.\n\n'
        '2020+: Production use. NOAA STOFS-3D-Atlantic (operational Jan 2023). '
        'EPA Chesapeake Bay Phase 7 model (2022). NextGen BMI integration (2023-present). '
        'Latest release: v5.11.0 (Feb 2025).'
    )

    doc.add_heading('Who Uses SCHISM?', level=3)
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    add_table_row(t, ['Organization', 'Use Case', 'Scale'], header=True)
    users = [
        ('NOAA/NOS', 'STOFS-3D-Atlantic operational forecasts', '2.9M nodes, 5.6M elements'),
        ('NOAA/OWP', 'NextGen Water Resources Framework', 'Experimental BMI integration'),
        ('US EPA', 'Chesapeake Bay Phase 7 water quality', 'Replaces 30+ year old Bay Model'),
        ('Oregon DOGAMI', 'Official tsunami inundation maps', 'NTHMP-certified'),
        ('Taiwan CWB', 'Daily ROCFORS operational forecasts', 'National coverage'),
        ('HZG Germany', 'North Sea & Baltic Sea dynamics', 'EU operational'),
        ('CA Dept. Water Resources', 'Bay-Delta water quality', 'State-level policy'),
    ]
    for u in users:
        add_table_row(t, u, bold_first=True)

    doc.add_heading('SCHISM by the Numbers', level=3)
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    add_table_row(t, ['Metric', 'Value'], header=True)
    stats = [
        ('Source files', '437+ Fortran files'),
        ('GitHub stars', '125+'),
        ('GitHub forks', '106+'),
        ('Contributors', '43+'),
        ('Total commits', '2,086+'),
        ('Latest release', 'v5.11.0 (Feb 2025)'),
        ('License', 'Apache 2.0'),
        ('Languages', 'Fortran 90/2003, C, Python'),
    ]
    for s in stats:
        add_table_row(t, s, bold_first=True)

    doc.add_page_break()

    # --- Section 2: Physics ---
    doc.add_heading('2. SCHISM Physics & What It Computes', level=2)

    doc.add_heading('Core Equations', level=3)
    doc.add_paragraph(
        'SCHISM solves the 3D Reynolds-averaged Navier-Stokes equations in hydrostatic form:'
    )

    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    add_table_row(t, ['Equation', 'What It Describes', 'ML Analogy'], header=True)
    eqs = [
        ('Continuity', 'Conservation of water mass', 'Batch normalization — conserving total'),
        ('Momentum (x)', 'East-west water acceleration', 'Horizontal gradient update'),
        ('Momentum (y)', 'North-south water acceleration', 'Vertical gradient update'),
        ('Vertical velocity', 'Up-down water movement', 'Skip connections'),
        ('Transport', 'Movement of heat, salt, tracers', 'Feature propagation through layers'),
        ('Equation of State', 'Density from T + S', 'Activation function'),
    ]
    for e in eqs:
        add_table_row(t, e, bold_first=True)

    doc.add_heading('Physical Processes Handled', level=3)
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    add_table_row(t, ['Process', 'Description', 'When It Matters'], header=True)
    processes = [
        ('Tides', 'Gravitational pull of moon/sun', 'Always (oceanic domains)'),
        ('Storm surge', 'Wind/pressure pushing water onshore', 'Hurricanes, nor\'easters'),
        ('River plumes', 'Freshwater spreading into saltwater', 'Estuaries, deltas'),
        ('Thermal stratification', 'Warm water on top, cold below', 'Lakes, deep water'),
        ('Salinity intrusion', 'Saltwater moving upstream', 'Droughts, sea level rise'),
        ('Wind-driven currents', 'Surface currents from wind stress', 'Open ocean, large lakes'),
        ('Coriolis effect', 'Earth rotation deflecting currents', 'Large domains'),
        ('Wetting/drying', 'Water flooding/receding from land', 'Coastal flooding, marshes'),
        ('Vegetation drag', 'Plants slowing water flow', 'Wetlands, marshes'),
        ('Waves', 'Wind-generated surface waves', 'Coupled with WWM-III'),
    ]
    for p in processes:
        add_table_row(t, p, bold_first=True)

    doc.add_heading('Inputs and Outputs', level=3)
    doc.add_paragraph('SCHISM Inputs (Forcing): Wind speed & direction, atmospheric pressure, '
                      'air temperature, humidity, rainfall, river discharge, ocean boundary tides, bathymetry.')
    doc.add_paragraph('SCHISM Outputs (Predictions): Water surface level (ETA2), current velocity (u, v, w), '
                      'water temperature, salinity, turbulent mixing, wetting/drying, wave parameters, '
                      'sediment transport, water quality.')

    doc.add_page_break()

    # --- Section 3: Unstructured Mesh ---
    doc.add_heading('3. The Unstructured Mesh', level=2)

    doc.add_paragraph(
        'SCHISM uses an unstructured mesh — a collection of triangles and quadrilaterals that tile the domain. '
        'Unlike a regular grid (same spacing everywhere), an unstructured mesh allows variable resolution. '
        'Fine resolution near complex coastlines (8m!) and coarse resolution in open ocean (2km).'
    )

    doc.add_heading('Triangles vs Quads', level=3)
    doc.add_paragraph(
        'SCHISM uniquely supports BOTH triangle (i34=3) and quadrilateral (i34=4) elements in the same mesh. '
        'Triangles are better for complex coastlines and irregular boundaries. '
        'Quads are better for channels, rivers, and straight boundaries — offering higher accuracy per element.'
    )

    doc.add_heading('Mesh Terminology', level=3)
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    add_table_row(t, ['Term', 'Definition', 'ML Analogy'], header=True)
    terms = [
        ('Node', 'A point/vertex in the mesh', 'Graph node in GNN'),
        ('Edge/Side', 'Line connecting two nodes', 'Edge in graph'),
        ('Element/Face', 'Triangle or quad (3-4 nodes)', '"Pixel" equivalent'),
        ('i34', 'Element type array (3=tri, 4=quad)', 'Node degree/type'),
        ('elnode', 'Which nodes form each element', 'Adjacency list'),
        ('isidenode', 'Which nodes form each edge', 'Edge index'),
    ]
    for term in terms:
        add_table_row(t, term, bold_first=True)

    doc.add_heading('Variable Staggering', level=3)
    doc.add_paragraph(
        'SCHISM uses Arakawa-like staggering:\n'
        '• Node (vertex): Water surface elevation (η), horizontal coordinates, depth\n'
        '• Side center (edge midpoint): Horizontal velocities (u, v)\n'
        '• Element center: Vertical velocity (w), tracer concentrations (T, S)'
    )

    doc.add_heading('Real-World Mesh Scales (STOFS-3D-Atlantic)', level=3)
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    add_table_row(t, ['Feature', 'Resolution'], header=True)
    scales = [
        ('Shoreline', '1.5-2 km'),
        ('Floodplain', '600 m'),
        ('Watershed rivers', '8 m (!)'),
        ('Levees', '2-10 m'),
        ('Total nodes', '2,926,236'),
        ('Total elements', '5,654,157'),
    ]
    for s in scales:
        add_table_row(t, s, bold_first=True)

    doc.add_page_break()

    # --- Section 4: Vertical Coordinates ---
    doc.add_heading('4. Vertical Coordinate System', level=2)
    doc.add_paragraph(
        'Ocean depth varies from 1m in marshes to 5,000m in deep ocean. SCHISM offers two vertical grid options:'
    )
    doc.add_paragraph(
        'Option A: SZ Hybrid (Sigma + Z-levels) — Sigma (terrain-following) layers near the surface, '
        'Z-levels (fixed depth) below a demarcation depth h_s. Control parameters theta_b and theta_f '
        'regulate vertical resolution distribution.'
    )
    doc.add_paragraph(
        'Option B: LSC2 (Localized Sigma Coordinates with Shaved Cell) — UNIQUE TO SCHISM! '
        'Each grid node has its OWN vertical coordinate set for maximum flexibility. '
        'No other ocean model has this capability. ML Analogy: Like adaptive computation in neural networks — '
        'more layers where needed (deep water), fewer where not (shallow water).'
    )

    # --- Section 5: Semi-Implicit ---
    doc.add_heading('5. Semi-Implicit Time Stepping', level=2)
    doc.add_paragraph(
        'Most ocean models use explicit time stepping with a CFL stability constraint: '
        'CFL = (velocity × dt) / dx must be < 1.0. For fine mesh (dx=10m) with fast flow (v=2 m/s), '
        'dt must be < 5 seconds. This is very expensive.'
    )
    doc.add_paragraph(
        'SCHISM uses semi-implicit stepping with NO CFL constraint. The same fine mesh can use '
        'dt = 120-300 seconds. This means 10-100x fewer time steps than explicit models.'
    )
    doc.add_paragraph(
        'COUNTERINTUITIVE: SCHISM actually WANTS large CFL numbers. When CFL < 0.4, numerical '
        'diffusion INCREASES. Large time steps = less diffusion = better accuracy.'
    )
    doc.add_paragraph(
        'The Eulerian-Lagrangian Method (ELM) for advection: Step 1 — Backtrack from current position '
        'along the flow to find where the water "came from" (Foot of Characteristic Line). '
        'Step 2 — Interpolate the value at that origin point. This method is unconditionally stable.'
    )

    doc.add_page_break()

    # --- Section 6: Comparison ---
    doc.add_heading('6. SCHISM vs Other Ocean Models', level=2)

    t = doc.add_table(rows=1, cols=5)
    t.style = 'Table Grid'
    add_table_row(t, ['Feature', 'SCHISM', 'ADCIRC', 'FVCOM', 'ROMS'], header=True)
    comp = [
        ('Grid Type', 'Hybrid tri+quad', 'Triangles only', 'Triangles only', 'Rectangular'),
        ('Time Stepping', 'Semi-implicit', 'Explicit', 'Explicit', 'Mode-split'),
        ('CFL Constraint', 'NONE', 'Required', 'Required', 'Required'),
        ('Max dt (fine mesh)', '120-300s', '1-10s', '1-10s', '10-60s'),
        ('Vertical Coords', 'LSC2/SZ (unique!)', 'Sigma only', 'Sigma', 'Terrain-following'),
        ('Model Polymorphism', '1D+2D+3D', '2D or 3D', '3D only', '3D only'),
        ('Wetting/Drying', 'Natural', 'Special treatment', 'Supported', 'Limited'),
        ('Cross-Scale', '8m → 2km', 'Limited', 'Limited', 'Fixed resolution'),
        ('BMI Support', 'Yes (LynkerIntel)', 'No', 'No', 'No'),
        ('NOAA Operational', 'STOFS-3D', 'ETSS, ESTOFS', 'NGOFS2', 'CBOFS, DBOFS'),
    ]
    for c in comp:
        add_table_row(t, c, bold_first=True)

    # --- Section 7: NOAA ---
    doc.add_heading('7. SCHISM in NOAA Operations', level=2)

    doc.add_heading('STOFS-3D-Atlantic', level=3)
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    add_table_row(t, ['Aspect', 'Details'], header=True)
    stofs = [
        ('Operational since', 'January 2023'),
        ('Run schedule', 'Daily at 12 UTC'),
        ('Forecast', '24-hour nowcast + up to 96-hour forecast'),
        ('Outputs', 'Water level, 2D/3D temperature, salinity, currents'),
        ('Grid', '2.9M nodes, 5.6M elements'),
        ('Resolution', '8m (rivers) to 2km (shoreline)'),
        ('Hydrology', 'Uses National Water Model (NWM) outputs'),
        ('Developed by', 'NOAA/NOS + NWS/NCEP + VIMS jointly'),
    ]
    for s in stofs:
        add_table_row(t, s, bold_first=True)

    doc.add_heading('NextGen Framework Integration', level=3)
    doc.add_paragraph(
        'GitHub Issue: NOAA-OWP/ngen#547 — "Evaluating SCHISM BMI as a NextGen Formulation". '
        'Status: Open (experimental, in development). '
        'Coupling: NWM (inland) → T-Route (routing) → SCHISM (coastal) via BMI. '
        'Key people: Jason Ducker (NOAA/NWS), Phil Miller, Y. Joseph Zhang (VIMS).'
    )

    doc.add_page_break()

    # --- Section 8: Modules ---
    doc.add_heading('8. SCHISM Modules & Capabilities', level=2)

    doc.add_heading('12 Tracer Modules', level=3)
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    add_table_row(t, ['#', 'Module', 'What It Computes'], header=True)
    mods = [
        ('1', 'TEM — Temperature', 'Water temperature distribution'),
        ('2', 'SAL — Salinity', 'Salt concentration'),
        ('3', 'GEN — Generic Tracer', 'User-customizable passive tracer'),
        ('4', 'AGE — Water Age', 'How long water has been in domain'),
        ('5', 'SED3D — 3D Sediment', 'Non-cohesive sediment transport'),
        ('6', 'EcoSim — Ecological Simulation', 'Marine ecosystem'),
        ('7', 'ICM — CE-QUAL-ICM', 'USACE water quality model'),
        ('8', 'CoSINE — C-Si-N Ecosystem', 'Carbon, Silicate, Nitrogen'),
        ('9', 'FIB — Fecal Indicator Bacteria', 'Bacteria tracking'),
        ('10', 'TIMOR', 'Currently inactive'),
        ('11', 'FABM — Aquatic BGC Framework', 'Generic biogeochemistry'),
        ('12', 'DVD — Numerical Mixing', 'Mixing analysis diagnostic'),
    ]
    for m in mods:
        add_table_row(t, m)

    doc.add_page_break()

    # =========================================================================
    # PART II: SCHISM BMI
    # =========================================================================
    doc.add_heading('PART II: SCHISM BMI — THE WRAPPER', level=1)

    # --- Section 10: Architecture ---
    doc.add_heading('10. SCHISM BMI Architecture', level=2)

    doc.add_heading('Who Built It?', level=3)
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    add_table_row(t, ['Aspect', 'Details'], header=True)
    arch = [
        ('Organization', 'LynkerIntel (Lynker Intelligence)'),
        ('Purpose', 'Enable SCHISM in NOAA NextGen framework'),
        ('Repository', 'github.com/LynkerIntel/SCHISM_BMI'),
        ('Branch', 'bmi-integration-master'),
        ('Main file', 'src/BMI/bmischism.f90 (1,729 lines)'),
        ('License', 'Apache 2.0'),
        ('Commits', '1,540+'),
    ]
    for a in arch:
        add_table_row(t, a, bold_first=True)

    doc.add_heading('Key Design Decision: Global State', level=3)
    doc.add_paragraph(
        'SCHISM is a massive legacy Fortran model (437 files). The BMI wrapper uses a TINY container '
        '(schism_type, 51 lines) for config/time tracking, but reads/writes DIRECTLY to schism_glbl '
        'module variables for all physics state. This is the non-invasive approach — wrap without modifying.'
    )

    # --- Section 11-13: Control ---
    doc.add_heading('11. BMI Initialize — How SCHISM Starts', level=2)
    doc.add_paragraph(
        'Step 1: Read BMI config file (Fortran namelist format) with model_start_time, model_end_time, '
        'time_step_size, SCHISM_dir, SCHISM_NSCRIBES.\n'
        'Step 2: Compute time parameters (num_time_steps = (end-start)/step_size).\n'
        'Step 3: Initialize MPI (call parallel_init with given_communicator).\n'
        'Step 4: Call SCHISM\'s own init (schism_init reads param.nml, hgrid.gr3, vgrid.in, '
        'allocates all global arrays — 7,074 lines of initialization!).\n'
        'Step 5: Store iths and ntime in container.\n\n'
        'IMPORTANT: MPI communicator must be set via set_value("bmi_mpi_comm_handle") BEFORE calling initialize().'
    )

    doc.add_heading('12. BMI Update — How SCHISM Steps', level=2)
    doc.add_paragraph(
        'update(): Increments iths by 1, then calls schism_step(iths) — 10,742 lines of physics. '
        'Solves continuity, momentum, transport equations. Updates eta2, uu2, vv2, etc.\n\n'
        'update_until(time): Loops calling update() until model_time >= target_time.'
    )

    doc.add_heading('13. BMI Finalize — How SCHISM Stops', level=2)
    doc.add_paragraph(
        'Calls schism_finalize (close output files, deallocate global arrays, 155 lines) '
        'then parallel_finalize (MPI_Finalize, shut down all MPI communications).'
    )

    doc.add_page_break()

    # --- Section 14: Input Variables ---
    doc.add_heading('14. BMI Input Variables (12 Detailed)', level=2)

    t = doc.add_table(rows=1, cols=7)
    t.style = 'Table Grid'
    add_table_row(t, ['#', 'Variable', 'Description', 'Units', 'Grid', 'Location', 'Type'], header=True)
    inputs = [
        ('1', 'Q_bnd_source', 'River discharge into domain', 'm³/s', '4 (SOURCE)', 'face', 'double'),
        ('2', 'Q_bnd_sink', 'Water extraction from domain', 'm³/s', '5 (SINK)', 'face', 'double'),
        ('3', 'ETA2_bnd', 'Water levels at open ocean boundary', 'm', '3 (OFFSHORE)', 'node', 'double'),
        ('4', 'SFCPRS', 'Surface atmospheric pressure', 'Pa', '1 (ALL_NODES)', 'node', 'double'),
        ('5', 'TMP2m', 'Air temperature at 2m', 'K', '1 (ALL_NODES)', 'node', 'double'),
        ('6', 'U10m', 'Eastward wind at 10m', 'm/s', '1 (ALL_NODES)', 'node', 'double'),
        ('7', 'V10m', 'Northward wind at 10m', 'm/s', '1 (ALL_NODES)', 'node', 'double'),
        ('8', 'SPFH2m', 'Specific humidity at 2m', 'kg/kg', '1 (ALL_NODES)', 'node', 'double'),
        ('9', 'RAINRATE', 'Precipitation rate (ADDITIVE!)', 'kg/m²/s', '2 (ALL_ELEM)', 'face', 'double'),
        ('10', 'ETA2_dt', 'Water level boundary update interval', 's', '7 (SCALAR)', '-', 'double'),
        ('11', 'Q_dt', 'Discharge update interval', 's', '8 (SCALAR)', '-', 'double'),
        ('12', 'bmi_mpi_comm_handle', 'MPI communicator handle', '-', '9 (SCALAR)', '-', 'integer'),
    ]
    for inp in inputs:
        add_table_row(t, inp)

    doc.add_heading('Input Variables by Category', level=3)
    doc.add_paragraph(
        'Atmospheric Forcing (5 vars): SFCPRS, TMP2m, U10m, V10m, SPFH2m — applied at every mesh node.\n'
        'Boundary Conditions (3 vars): ETA2_bnd, Q_bnd_source, Q_bnd_sink — applied at domain edges.\n'
        'Precipitation (1 var): RAINRATE — applied at every mesh element (ADDS to existing source!).\n'
        'Time Control (2 vars): ETA2_dt, Q_dt — scalar values controlling update frequency.\n'
        'System (1 var): bmi_mpi_comm_handle — MPI communicator (must be set before initialize!).'
    )

    doc.add_page_break()

    # --- Section 15: Output Variables ---
    doc.add_heading('15. BMI Output Variables (5 Detailed)', level=2)

    t = doc.add_table(rows=1, cols=6)
    t.style = 'Table Grid'
    add_table_row(t, ['#', 'Variable', 'Description', 'Units', 'Grid', 'Source Array'], header=True)
    outputs = [
        ('1', 'ETA2', 'Water surface elevation', 'm', '1 (ALL_NODES)', 'eta2(:)'),
        ('2', 'VX', 'Eastward current velocity (surface)', 'm/s', '1 (ALL_NODES)', 'uu2(1,:)'),
        ('3', 'VY', 'Northward current velocity (surface)', 'm/s', '1 (ALL_NODES)', 'vv2(1,:)'),
        ('4', 'TROUTE_ETA2', 'Water levels at T-Route stations', 'm', '6 (STATION)', 'sta_out_gb(:,1)'),
        ('5', 'BEDLEVEL', 'Bed elevation (INVERTED from dp)', 'm', '1 (ALL_NODES)', '-1.0 * dp(:)'),
    ]
    for o in outputs:
        add_table_row(t, o)

    doc.add_paragraph(
        '\nETA2 is THE primary output — water surface height relative to datum. '
        'This is what WRF-Hydro would receive for 2-way coupling.\n\n'
        'VX/VY expose only surface layer (index 1) of the 3D velocity field.\n\n'
        'BEDLEVEL is INVERTED: SCHISM stores depth as positive below datum (dp=10 means 10m below), '
        'but BMI returns BEDLEVEL = -dp (negative below datum, positive above).'
    )

    doc.add_page_break()

    # --- Section 16: Grid System ---
    doc.add_heading('16. BMI Grid System (9 Grids)', level=2)

    t = doc.add_table(rows=1, cols=5)
    t.style = 'Table Grid'
    add_table_row(t, ['Grid ID', 'Constant Name', 'Type', 'Rank', 'Variables'], header=True)
    grids = [
        ('1', 'GRID_ALL_NODES', 'unstructured', '2', 'ETA2, VX, VY, atmos forcing, BEDLEVEL'),
        ('2', 'GRID_ALL_ELEMENTS', 'points', '2', 'RAINRATE'),
        ('3', 'GRID_OFFSHORE_BOUNDARY', 'points', '2', 'ETA2_bnd'),
        ('4', 'GRID_SOURCE_ELEMENTS', 'points', '1', 'Q_bnd_source'),
        ('5', 'GRID_SINK_ELEMENTS', 'points', '1', 'Q_bnd_sink'),
        ('6', 'GRID_STATION_POINTS', 'points', '2', 'TROUTE_ETA2'),
        ('7', 'ETA2_TIMESTEP', 'scalar', '1', 'ETA2_dt'),
        ('8', 'Q_TIMESTEP', 'scalar', '1', 'Q_dt'),
        ('9', 'MPI_COMMUNICATOR', 'scalar', '1', 'bmi_mpi_comm_handle'),
    ]
    for g in grids:
        add_table_row(t, g, bold_first=True)

    doc.add_heading('Grid Types Explained', level=3)
    doc.add_paragraph(
        '"unstructured" (Grid 1 only): Full triangle/quad mesh with complete node/edge/face topology. '
        'Has get_grid_x/y/z, node/edge/face_count, edge_nodes, face_nodes, face_edges, nodes_per_face. '
        'Does NOT have shape, spacing, origin (not applicable for unstructured mesh!).\n\n'
        '"points" (Grids 2-6): Collection of x,y,z coordinates with NO connectivity. Just scatter points.\n\n'
        '"scalar" (Grids 7-9): Single value, no spatial extent. size=1, rank=1.'
    )

    doc.add_page_break()

    # --- Section 17: t0/t1 Pattern ---
    doc.add_heading('17. The t0/t1 Sliding Window Pattern', level=2)

    doc.add_paragraph(
        'SCHISM stores TWO time snapshots for every forcing variable — "previous" (t0) and "current" (t1). '
        'When new data arrives via set_value(), the old t1 slides to t0 and new data goes into t1.'
    )

    doc.add_heading('Complete t0/t1 Variable Pairs', level=3)
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    add_table_row(t, ['BMI Variable', 't0 Array', 't1 Array', 'Description'], header=True)
    pairs = [
        ('SFCPRS', 'pr1(:)', 'pr2(:)', 'Surface pressure'),
        ('TMP2m', 'airt1(:)', 'airt2(:)', 'Air temperature'),
        ('U10m', 'windx1(:)', 'windx2(:)', 'Eastward wind'),
        ('V10m', 'windy1(:)', 'windy2(:)', 'Northward wind'),
        ('SPFH2m', 'shum1(:)', 'shum2(:)', 'Specific humidity'),
        ('ETA2_bnd', 'ath2(:,:,:,1,:)', 'ath2(:,:,:,2,:)', 'Boundary water levels'),
        ('Q_bnd_source', 'ath3(:,1,1,1)', 'ath3(:,1,2,1)', 'Source discharge'),
        ('Q_bnd_sink', 'ath3(:,1,1,2)', 'ath3(:,1,2,2)', 'Sink discharge'),
    ]
    for p in pairs:
        add_table_row(t, p, bold_first=True)

    doc.add_heading('Why Two Time Slots?', level=3)
    doc.add_paragraph(
        'SCHISM interpolates between t0 and t1 during each timestep: '
        'forcing(t) = t0_value + (t - t0_time)/(t1_time - t0_time) × (t1_value - t0_value). '
        'This creates smooth forcing transitions, avoiding numerical instabilities from abrupt jumps.\n\n'
        'ML Analogy: Like a linear learning rate schedule — smooth transitions between values, '
        'not sudden jumps.'
    )

    # --- Section 18: RAINRATE ---
    doc.add_heading('18. RAINRATE — The Special Variable', level=2)
    doc.add_paragraph(
        'Every other input uses the standard t0/t1 slide-and-replace. RAINRATE is the ONLY variable '
        'that ADDS to existing values instead of replacing them.\n\n'
        'Reason: Q_bnd_source is set first (river discharge into ath3), then RAINRATE adds rain '
        'contribution ON TOP. Both contribute to the same source term (ath3).\n\n'
        'Unit conversion: Q_rain = RAINRATE × element_area / 1000\n'
        '(kg/m²/s × m² / kg/m³ = m³/s)\n\n'
        'ML Analogy: Like a residual connection — output = existing_value + new_contribution.'
    )

    doc.add_page_break()

    # --- Section 19: Variable Info ---
    doc.add_heading('19. Variable Info Functions', level=2)

    t = doc.add_table(rows=1, cols=6)
    t.style = 'Table Grid'
    add_table_row(t, ['Variable', 'Type', 'Units', 'Grid', 'Itemsize', 'Location'], header=True)
    all_vars = [
        ('ETA2', 'double precision', 'm', '1', '8 bytes', 'node'),
        ('VX', 'double precision', 'm s-1', '1', '8 bytes', 'node'),
        ('VY', 'double precision', 'm s-1', '1', '8 bytes', 'node'),
        ('TROUTE_ETA2', 'double precision', 'm', '6', '8 bytes', 'node'),
        ('BEDLEVEL', 'double precision', 'm', '1', '8 bytes', 'node'),
        ('Q_bnd_source', 'double precision', 'm3 s-1', '4', '8 bytes', 'face'),
        ('Q_bnd_sink', 'double precision', 'm3 s-1', '5', '8 bytes', 'face'),
        ('ETA2_bnd', 'double precision', 'm', '3', '8 bytes', 'node'),
        ('SFCPRS', 'double precision', 'Pa', '1', '8 bytes', 'node'),
        ('TMP2m', 'double precision', 'K', '1', '8 bytes', 'node'),
        ('U10m', 'double precision', 'm s-1', '1', '8 bytes', 'node'),
        ('V10m', 'double precision', 'm s-1', '1', '8 bytes', 'node'),
        ('SPFH2m', 'double precision', 'kg kg-1', '1', '8 bytes', 'node'),
        ('RAINRATE', 'double precision', 'kg m-2 s-1', '2', '8 bytes', 'face'),
        ('ETA2_dt', 'double precision', 's', '7', '8 bytes', '-'),
        ('Q_dt', 'double precision', 's', '8', '8 bytes', '-'),
        ('bmi_mpi_comm', 'integer', '-', '9', '4 bytes', '-'),
    ]
    for v in all_vars:
        add_table_row(t, v)

    # --- Section 20-22 ---
    doc.add_heading('20. Get/Set Value Patterns', level=2)
    doc.add_paragraph(
        'get_value: Uses select case to map variable name to global array. '
        'Output arrays are already 1D (no reshape needed unlike Heat model). '
        'ETA2 → eta2(:), VX → uu2(1,:), VY → vv2(1,:), BEDLEVEL → -1.0*dp(:).\n\n'
        'set_value: Standard pattern applies t0/t1 sliding (var_t0 = var_t1; var_t1 = new_data). '
        'Exception: RAINRATE adds instead of replacing.\n\n'
        'get_value_at_indices: Reads only specific node indices (efficient for partial queries). '
        'Available for ETA2, TROUTE_ETA2, VX, VY, BEDLEVEL.\n\n'
        'get_value_ptr: Returns BMI_FAILURE for ALL variables (not implemented).\n'
        'set_value_float: Returns BMI_FAILURE (all SCHISM vars are double precision).'
    )

    doc.add_heading('21. Time Functions', level=2)
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    add_table_row(t, ['Function', 'Returns', 'Source'], header=True)
    times = [
        ('get_start_time()', 'model_start_time', 'From config namelist'),
        ('get_end_time()', 'model_end_time', 'From config or computed'),
        ('get_current_time()', 'iths × dt', 'Computed at runtime'),
        ('get_time_step()', 'time_step_size', 'From config namelist'),
        ('get_time_units()', '"s"', 'Hardcoded (always seconds)'),
    ]
    for ti in times:
        add_table_row(t, ti, bold_first=True)

    doc.add_page_break()

    # --- Section 22-24 ---
    doc.add_heading('22. NextGen Integration', level=2)
    doc.add_paragraph(
        'Conditional compilation flags:\n'
        '• #ifdef NGEN_ACTIVE: Uses bmif_2_0_iso (ISO C binding), adds register_bmi() function. '
        'Required for NextGen framework discovery.\n'
        '• #ifdef OLDIO: Serial I/O mode (each rank dumps its own data). Used for serial BMI.\n\n'
        'register_bmi(): Allocates a bmi_schism instance, wraps in C pointer via "box" wrapper, '
        'returns pointer to caller. NextGen calls this C function to discover SCHISM.\n\n'
        'iso_c_fortran_bmi/: Provides the C interop bridge (iso_c_bmi.f90, 39.8 KB) — wraps every '
        'BMI function with ISO_C_BINDING for C/Python interoperability.'
    )

    doc.add_heading('23. Build System & Configuration', level=2)
    doc.add_heading('CMake Build Flags', level=3)
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    add_table_row(t, ['Flag', 'Default', 'Purpose'], header=True)
    flags = [
        ('USE_BMI', 'ON', 'Enable BMI wrapper compilation'),
        ('BUILD_SHARED_LIBS', 'ON', 'Build libschism_bmi.so'),
        ('OLDIO', 'ON', 'Serial I/O mode'),
        ('USE_ATMOS', 'ON', 'Atmospheric forcing via BMI'),
        ('USE_GOTM', 'OFF', 'GOTM turbulence model'),
        ('USE_WWM', 'OFF', 'Wind wave model'),
        ('USE_ICE', 'OFF', 'Sea ice model'),
        ('USE_SED', 'OFF', 'Sediment transport'),
    ]
    for f in flags:
        add_table_row(t, f, bold_first=True)

    doc.add_heading('24. Current Limitations', level=2)
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    add_table_row(t, ['#', 'Limitation', 'Impact'], header=True)
    lims = [
        ('1', 'MPI partially implemented', 'Prefer serial mode for BMI'),
        ('2', 'Connectors not fully validated', 'Use with caution'),
        ('3', 'get_value_ptr not implemented', 'No zero-copy access'),
        ('4', 'set_value_float not implemented', 'All vars are double precision'),
        ('5', 'No CSDMS Standard Names', 'Uses internal names (ETA2, VX)'),
        ('6', 'Only surface velocity exposed', 'VX/VY = surface only, not full 3D'),
        ('7', 'No PyMT pathway', 'Babelizer not run, no pymt_schism'),
        ('8', 'Not in CSDMS catalog', 'Built for NextGen, not PyMT'),
    ]
    for l in lims:
        add_table_row(t, l)

    doc.add_page_break()

    # --- Section 25: Links ---
    doc.add_heading('25. Repository Links & References', level=2)

    doc.add_heading('Official Repositories', level=3)
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    add_table_row(t, ['Resource', 'URL'], header=True)
    links = [
        ('SCHISM Official Website', 'http://ccrm.vims.edu/schismweb/'),
        ('SCHISM GitHub (main)', 'https://github.com/schism-dev/schism'),
        ('SCHISM Online Docs', 'https://schism-dev.github.io/schism/master/index.html'),
        ('SCHISM Wiki', 'http://ccrm.vims.edu/w/index.php/About_SCHISM'),
        ('SCHISM BMI (LynkerIntel)', 'https://github.com/LynkerIntel/SCHISM_BMI'),
        ('SCHISM NWM BMI (schism-dev)', 'https://github.com/schism-dev/schism_NWM_BMI'),
        ('PySchism', 'https://github.com/schism-dev/pyschism'),
        ('SCHISM-ESMF', 'https://github.com/schism-dev/schism-esmf'),
        ('NOAA NextGen (ngen)', 'https://github.com/NOAA-OWP/ngen'),
        ('NextGen SCHISM Issue', 'https://github.com/NOAA-OWP/ngen/issues/547'),
        ('UFS Coastal App', 'https://github.com/oceanmodeling/ufs-coastal-app'),
        ('NOAA STOFS-3D Data', 'https://registry.opendata.aws/noaa-nos-stofs3d/'),
        ('SCHISM v5.8 Manual PDF', 'https://ccrm.vims.edu/schismweb/SCHISM_v5.8-Manual.pdf'),
        ('SCHISM 2016 Paper PDF', 'https://ccrm.vims.edu/yinglong/Courses/Marsh-2017/Zhang_etal_OM_2016-SCHISMpaper.pdf'),
        ('Dr. Zhang VIMS Profile', 'https://www.vims.edu/about/directory/faculty/zhang_yj.php'),
    ]
    for l in links:
        add_table_row(t, l, bold_first=True)

    doc.add_heading('Key Publications', level=3)
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    add_table_row(t, ['Year', 'Citation', 'Topic'], header=True)
    pubs = [
        ('2008', 'Zhang & Baptista, Ocean Modelling 21(3-4), 71-96', 'Original SELFE paper'),
        ('2015', 'Zhang et al., Ocean Modelling 85, 16-31', 'LSC2 vertical coordinates'),
        ('2016', 'Zhang, Ye, Stanev, Grashorn, Ocean Modelling 102, 64-81', 'THE SCHISM paper'),
        ('2020', 'Zhang et al., Ocean Dynamics 70, 621-640', 'Compound flooding'),
    ]
    for p in pubs:
        add_table_row(t, p, bold_first=True)

    doc.add_page_break()

    # --- Section 26: Summary ---
    doc.add_heading('26. Summary & Key Takeaways', level=2)

    doc.add_heading('SCHISM Model Summary', level=3)
    doc.add_paragraph(
        '• Grid: Unstructured triangles + quads (variable resolution, 8m to 2km)\n'
        '• Time: Semi-implicit (NO CFL constraint, large dt OK)\n'
        '• Vertical: LSC2 (each node has its own vertical levels — unique!)\n'
        '• Advection: ELM (unconditionally stable)\n'
        '• Scale: Creek-to-ocean in one simulation\n'
        '• Modules: 12 tracer + 10 non-tracer\n'
        '• Operations: NOAA STOFS-3D-Atlantic (since Jan 2023)\n'
        '• Size: 437 files, 100K+ lines of Fortran\n'
        '• License: Apache 2.0'
    )

    doc.add_heading('SCHISM BMI Summary', level=3)
    doc.add_paragraph(
        '• Built by: LynkerIntel (for NOAA NextGen)\n'
        '• Main file: bmischism.f90 (1,729 lines)\n'
        '• State: Global variables in schism_glbl (NOT embedded in wrapper type)\n'
        '• 12 Input Variables: 5 atmospheric + 3 boundary + 1 precip + 2 time + 1 MPI\n'
        '• 5 Output Variables: ETA2, VX, VY, TROUTE_ETA2, BEDLEVEL\n'
        '• 9 Grids: 1 unstructured + 5 points + 3 scalar\n'
        '• Key Pattern: t0/t1 sliding window for temporal interpolation\n'
        '• Special: RAINRATE adds instead of replaces (residual pattern)\n'
        '• Conditional compilation: #ifdef NGEN_ACTIVE, OLDIO'
    )

    doc.add_heading('Why This Matters for WRF-Hydro BMI', level=3)
    doc.add_paragraph(
        '1. Global state pattern → WRF-Hydro uses RT_DOMAIN (similar to schism_glbl)\n'
        '2. Config-only container → We\'ll create wrf_hydro_type for config/time only\n'
        '3. Delegate to model → We\'ll call land_driver_ini/exe, HYDRO_ini/exe\n'
        '4. select case dispatch → Same pattern for all variable functions\n'
        '5. Multiple grids → Grid 0 (1km), Grid 1 (250m), Grid 2 (network)\n'
        '6. Conditional compilation → #ifdef USE_NWM_BMI\n'
        '7. Non-invasive wrapper → Wrap WRF-Hydro without modifying it\n\n'
        'SCHISM BMI is our REAL-WORLD REFERENCE.\n'
        'Heat BMI is our TEMPLATE.\n'
        'Together they guide our bmi_wrf_hydro.f90 implementation.'
    )

    # --- Footer ---
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('— End of Document —')
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.size = Pt(10)

    return doc


if __name__ == '__main__':
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(script_dir, '13_SCHISM_And_Its_BMI_Complete_Deep_Dive.docx')

    print("Creating SCHISM & BMI Deep Dive Word Document...")
    doc = create_document()
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    print("Done!")
